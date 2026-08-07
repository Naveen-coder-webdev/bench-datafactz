"""
parsing.py — Stage 2: extract raw text from every 'pending' document.

Routing logic:
  pdf   -> try pdftotext first (digital PDFs). If it yields ~no text,
           fall back to OCR (pdf2image + tesseract) — this is how we
           distinguish digital vs scanned PDFs without being told which is which.
  docx  -> python-docx paragraph + table extraction
  xlsx  -> openpyxl, all sheets, all cells, row-major
  csv   -> plain read
  eml   -> stdlib email parser; body + attachment text (if text-decodable)
  txt/html -> plain read (html stripped of tags)
  png   -> tesseract OCR directly
"""

import os
import sqlite3
import subprocess
import email
from email import policy
import openpyxl
from docx import Document as DocxDocument
from html.parser import HTMLParser


class _HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.chunks = []

    def handle_data(self, data):
        self.chunks.append(data)

    def text(self):
        return " ".join(self.chunks)


def _pdftotext(path: str) -> str:
    try:
        result = subprocess.run(["pdftotext", "-layout", path, "-"],
                                 capture_output=True, text=True, timeout=30)
        return result.stdout or ""
    except Exception:
        return ""


def _ocr_pdf(path: str) -> str:
    try:
        from pdf2image import convert_from_path
        import pytesseract
        images = convert_from_path(path)
        text_parts = [pytesseract.image_to_string(img) for img in images]
        return "\n".join(text_parts)
    except Exception as e:
        return f"[OCR FAILED: {e}]"


def _ocr_image(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
        return pytesseract.image_to_string(Image.open(path))
    except Exception as e:
        return f"[OCR FAILED: {e}]"


def parse_pdf(path: str) -> tuple[str, str]:
    text = _pdftotext(path)
    if len(text.strip()) < 20:  # effectively no extractable text -> scanned PDF
        return _ocr_pdf(path), "tesseract_ocr"
    return text, "pdftotext"


def parse_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def parse_xlsx(path: str) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            parts.append(" | ".join(str(c) if c is not None else "" for c in row))
    return "\n".join(parts)


def parse_csv(path: str) -> str:
    with open(path, "r", errors="replace") as f:
        return f.read()


def parse_txt(path: str) -> str:
    with open(path, "r", errors="replace") as f:
        return f.read()


def parse_html(path: str) -> str:
    with open(path, "r", errors="replace") as f:
        raw = f.read()
    extractor = _HTMLTextExtractor()
    extractor.feed(raw)
    return extractor.text()


def parse_eml(path: str) -> str:
    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=policy.default)

    parts = [f"Subject: {msg.get('Subject', '')}", f"From: {msg.get('From', '')}", f"To: {msg.get('To', '')}"]

    body = msg.get_body(preferencelist=("plain",))
    if body:
        parts.append(body.get_content())

    for attachment in msg.iter_attachments():
        fname = attachment.get_filename() or "attachment"
        try:
            content = attachment.get_content()
            if isinstance(content, bytes):
                content = content.decode("utf-8", errors="replace")
            parts.append(f"[ATTACHMENT: {fname}]\n{content}")
        except Exception as e:
            parts.append(f"[ATTACHMENT: {fname} — could not decode: {e}]")

    return "\n".join(parts)


def parse_png(path: str) -> str:
    return _ocr_image(path)


PARSERS = {
    "docx": lambda p: (parse_docx(p), "docx_native"),
    "xlsx": lambda p: (parse_xlsx(p), "xlsx_native"),
    "csv": lambda p: (parse_csv(p), "text_native"),
    "txt": lambda p: (parse_txt(p), "text_native"),
    "html": lambda p: (parse_html(p), "html_strip"),
    "eml": lambda p: (parse_eml(p), "eml_native"),
    "png": lambda p: (parse_png(p), "tesseract_ocr"),
    "pdf": parse_pdf,
}


def parse_corpus(corpus_dir: str, db_path: str):
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute("SELECT doc_id, filename, file_type FROM documents WHERE status = 'pending'")
    rows = cur.fetchall()

    ok, failed = 0, 0
    for doc_id, filename, file_type in rows:
        path = os.path.join(corpus_dir, filename)
        parser = PARSERS.get(file_type)
        if not parser:
            cur.execute("UPDATE documents SET status='quarantined', quarantine_reason=? WHERE doc_id=?",
                        (f"No parser registered for file_type '{file_type}'", doc_id))
            failed += 1
            continue
        try:
            text, method = parser(path)
            cur.execute(
                "UPDATE documents SET status='parsed', raw_text=?, parse_method=?, parsed_at=CURRENT_TIMESTAMP WHERE doc_id=?",
                (text, method, doc_id),
            )
            ok += 1
        except Exception as e:
            cur.execute("UPDATE documents SET status='quarantined', quarantine_reason=? WHERE doc_id=?",
                        (f"Parse error: {e}", doc_id))
            failed += 1

    conn.commit()
    conn.close()
    print(f"Parsed {ok} documents, {failed} failed/quarantined during parsing")
    return ok, failed


if __name__ == "__main__":
    import sys
    corpus_dir = sys.argv[1] if len(sys.argv) > 1 else "../../corpus_generator/output/corpus"
    db_path = sys.argv[2] if len(sys.argv) > 2 else "../../db/breach.db"
    parse_corpus(corpus_dir, db_path)
