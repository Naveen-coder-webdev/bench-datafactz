"""
planters.py — one function per file type. Each planter takes person record(s),
writes a realistic file to disk, and logs exactly what it planted into the
Manifest. PII is embedded in natural context (prose, table cells, signature
blocks, image text) — never neatly labeled "SSN: ...".
"""

import os
import random
import textwrap
from email.message import EmailMessage
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.lib.units import inch

import openpyxl

OUT = "output/corpus"


# ----------------------------------------------------------------------
# Narrative helpers — realistic sentence templates PII gets embedded in
# ----------------------------------------------------------------------

LETTER_TEMPLATES = [
    "This letter confirms that {name} (DOB {dob}) is enrolled in our benefits "
    "program under policy holder SSN {ssn}. Please direct correspondence to "
    "{address} or by phone at {phone}.",
    "Re: Account Update for {name}\n\nWe have updated our records for {name}, "
    "residing at {address}. For verification purposes our system has your "
    "Social Security Number on file as {ssn} and driver's license {dl}.",
    "Dear {name},\n\nYour recent claim has been processed. Please confirm your "
    "date of birth ({dob}) and mailing address ({address}) are current. "
    "Reach us at {phone} or {email} with questions.",
]

INCIDENT_TEMPLATES = [
    "On the date of intake, patient {name} (DOB {dob}) presented with symptoms "
    "consistent with {condition}. Contact on file: {phone}, {email}. "
    "Insurance card number ending in the account on file.",
    "Employee {name}, job title {title} at {employer}, reported the incident. "
    "Employee ID references cross-check against SSN {ssn} were completed by HR. "
    "Home address on file: {address}.",
]

EMAIL_TEMPLATES = [
    "Hi team,\n\nCan you please update the file for {name}? Their new phone "
    "number is {phone} and email is {email}. Home address on record: "
    "{address}.\n\nThanks,\n{sender}",
    "Hello,\n\nAttached is the record for {name} (DOB {dob}). Please verify "
    "the SSN on file ({ssn}) matches the enrollment form before we proceed.\n\n"
    "Best,\n{sender}",
    "FYI — {name}'s card on file ({card}, exp {card_exp}) was declined during "
    "the last billing cycle. Please contact them at {phone} to update payment.\n\n"
    "{sender}",
]

FIRST_NAMES_STAFF = ["Alex Rivera", "Morgan Blake", "Jamie Chen", "Taylor Brooks"]


def _rand_signature():
    name = random.choice(FIRST_NAMES_STAFF)
    return f"{name}\n{random.choice(['HR Coordinator','Billing Specialist','Compliance Analyst','Account Manager'])}"


# ----------------------------------------------------------------------
# DOCX
# ----------------------------------------------------------------------

def plant_docx(doc_id, person, manifest, edge_case=None):
    filename = f"{doc_id}.docx"
    path = os.path.join(OUT, filename)
    doc = Document()

    template = random.choice(LETTER_TEMPLATES)
    name = random.choice(person["aliases"] + [person["full_name"]]) if edge_case == "name_variant" else person["full_name"]

    body = template.format(
        name=name, dob=person["dob"], ssn=person["ssn"], dl=person["dl_number"],
        address=person["address"], phone=person["phone"], email=person["email"],
    )

    doc.add_heading(random.choice(["Account Correspondence", "Policyholder Notice", "Member Services Letter"]), level=1)
    for para in body.split("\n\n"):
        p = doc.add_paragraph(para)
        p.style.font.size = Pt(11)
    doc.save(path)

    manifest.add_document(doc_id, filename, "docx")
    manifest.plant(doc_id, person["person_id"], "full_name", name, edge_case=edge_case)
    manifest.plant(doc_id, person["person_id"], "dob", person["dob"])
    manifest.plant(doc_id, person["person_id"], "ssn", person["ssn"])
    manifest.plant(doc_id, person["person_id"], "home_address", person["address"])
    manifest.plant(doc_id, person["person_id"], "phone", person["phone"])
    if "dl_number" in template:
        manifest.plant(doc_id, person["person_id"], "dl_number", person["dl_number"])
    return path


# ----------------------------------------------------------------------
# Digital PDF (text layer present — direct extraction, no OCR needed)
# ----------------------------------------------------------------------

def plant_digital_pdf(doc_id, person, manifest, edge_case=None):
    filename = f"{doc_id}.pdf"
    path = os.path.join(OUT, filename)
    c = pdfcanvas.Canvas(path, pagesize=letter)
    width, height = letter

    template = random.choice(INCIDENT_TEMPLATES)
    text = template.format(
        name=person["full_name"], dob=person["dob"], condition=person["medical_condition"],
        phone=person["phone"], email=person["email"], title=person["job_title"],
        employer=person["employer"], ssn=person["ssn"], address=person["address"],
    )

    c.setFont("Helvetica-Bold", 14)
    c.drawString(inch, height - inch, "Incident / Intake Report")
    c.setFont("Helvetica", 10)
    y = height - 1.4 * inch
    for line in textwrap.wrap(text, 90):
        c.drawString(inch, y, line)
        y -= 14
    c.save()

    manifest.add_document(doc_id, filename, "pdf_digital")
    manifest.plant(doc_id, person["person_id"], "full_name", person["full_name"], edge_case=edge_case)
    manifest.plant(doc_id, person["person_id"], "dob", person["dob"])
    if "medical" in template.lower() or "symptoms" in template:
        manifest.plant(doc_id, person["person_id"], "medical", person["medical_condition"])
    if "SSN" in template:
        manifest.plant(doc_id, person["person_id"], "ssn", person["ssn"])
    manifest.plant(doc_id, person["person_id"], "phone", person["phone"])
    manifest.plant(doc_id, person["person_id"], "email", person["email"])
    if "address" in template:
        manifest.plant(doc_id, person["person_id"], "home_address", person["address"])
    return path


# ----------------------------------------------------------------------
# Scanned / image-based PDF (NO text layer — requires OCR)
# ----------------------------------------------------------------------

def _get_font(size=22):
    try:
        return ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", size)
    except Exception:
        return ImageFont.load_default()


def plant_scanned_pdf(doc_id, person, manifest, edge_case=None):
    filename = f"{doc_id}.pdf"
    path = os.path.join(OUT, filename)
    img_path = os.path.join(OUT, f"_tmp_{doc_id}.png")

    img = Image.new("RGB", (1700, 2200), "white")
    draw = ImageDraw.Draw(img)
    font_title = _get_font(30)
    font_body = _get_font(22)

    draw.text((100, 80), "MEDICAL RECORDS DEPARTMENT — FAX COVER", font=font_title, fill="black")
    lines = [
        f"Patient Name: {person['full_name']}",
        f"Date of Birth: {person['dob']}",
        f"SSN on file: {person['ssn']}",
        f"Address: {person['address']}",
        f"Phone: {person['phone']}",
        f"Diagnosis notes: {person['medical_condition']}, ongoing management.",
        "This document was scanned from a physical fax and contains no digital text layer.",
    ]
    y = 220
    for line in lines:
        for wrapped in textwrap.wrap(line, 55):
            draw.text((100, y), wrapped, font=font_body, fill="black")
            y += 40
        y += 15

    # slight rotation + noise to look like a real scan
    img = img.rotate(random.uniform(-1.2, 1.2), fillcolor="white", expand=False)
    img.save(img_path)

    img_rgb = img.convert("RGB")
    img_rgb.save(path, "PDF", resolution=150.0)
    os.remove(img_path)

    manifest.add_document(doc_id, filename, "pdf_scanned", notes="No text layer — OCR required")
    manifest.plant(doc_id, person["person_id"], "full_name", person["full_name"], edge_case=edge_case or "requires_ocr")
    manifest.plant(doc_id, person["person_id"], "dob", person["dob"])
    manifest.plant(doc_id, person["person_id"], "ssn", person["ssn"])
    manifest.plant(doc_id, person["person_id"], "home_address", person["address"])
    manifest.plant(doc_id, person["person_id"], "phone", person["phone"])
    manifest.plant(doc_id, person["person_id"], "medical", person["medical_condition"])
    return path


# ----------------------------------------------------------------------
# PNG screenshot (e.g., "image of a spreadsheet" problem file)
# ----------------------------------------------------------------------

def plant_screenshot(doc_id, people_subset, manifest, edge_case="image_of_spreadsheet"):
    filename = f"{doc_id}.png"
    path = os.path.join(OUT, filename)
    rows = len(people_subset) + 1
    img = Image.new("RGB", (1400, 60 + rows * 40), "white")
    draw = ImageDraw.Draw(img)
    font = _get_font(18)
    headers = ["Name", "DOB", "SSN", "Phone"]
    x_positions = [40, 400, 650, 950]
    draw.rectangle([20, 20, 1380, 55], fill="#d9d9d9")
    for h, x in zip(headers, x_positions):
        draw.text((x, 28), h, font=font, fill="black")
    y = 60
    for p in people_subset:
        draw.text((x_positions[0], y + 8), p["full_name"], font=font, fill="black")
        draw.text((x_positions[1], y + 8), p["dob"], font=font, fill="black")
        draw.text((x_positions[2], y + 8), p["ssn"], font=font, fill="black")
        draw.text((x_positions[3], y + 8), p["phone"], font=font, fill="black")
        y += 40
    img.save(path)

    manifest.add_document(doc_id, filename, "png_screenshot", notes="Screenshot of tabular export — requires OCR + layout parsing")
    for p in people_subset:
        manifest.plant(doc_id, p["person_id"], "full_name", p["full_name"], edge_case=edge_case)
        manifest.plant(doc_id, p["person_id"], "dob", p["dob"])
        manifest.plant(doc_id, p["person_id"], "ssn", p["ssn"])
        manifest.plant(doc_id, p["person_id"], "phone", p["phone"])
    return path


# ----------------------------------------------------------------------
# XLSX (multi-person spreadsheet export)
# ----------------------------------------------------------------------

def plant_xlsx(doc_id, people_subset, manifest, edge_case=None, include_fp_traps=True):
    filename = f"{doc_id}.xlsx"
    path = os.path.join(OUT, filename)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Roster Export"

    headers = ["EmployeeID", "Full Name", "DOB", "SSN", "Phone", "Email", "Home Address", "Card on File"]
    ws.append(headers)

    for i, p in enumerate(people_subset):
        name = p["full_name"]
        # occasionally use a nickname/alias variant to force resolution work
        if edge_case == "name_variant" and p["aliases"] and random.random() < 0.4:
            name = random.choice(p["aliases"])
        ws.append([
            f"EMP-{1000+i}", name, p["dob"], p["ssn"], p["phone"], p["email"],
            p["address"], p["card_number"],
        ])
        manifest.plant(doc_id, p["person_id"], "full_name", name, edge_case=edge_case)
        manifest.plant(doc_id, p["person_id"], "dob", p["dob"])
        manifest.plant(doc_id, p["person_id"], "ssn", p["ssn"])
        manifest.plant(doc_id, p["person_id"], "phone", p["phone"])
        manifest.plant(doc_id, p["person_id"], "email", p["email"])
        manifest.plant(doc_id, p["person_id"], "home_address", p["address"])
        manifest.plant(doc_id, p["person_id"], "card_number", p["card_number"])

    if include_fp_traps:
        # order number that looks like an SSN pattern, and a placeholder/test row
        ws.append(["EMP-9999", "TEST USER", "1900-01-01", "000-00-0000", "555-000-0000", "test@example.com", "123 Placeholder St", "0000-0000-0000-0000"])
        ws.append(["ORD-7788", f"Order Ref {random.randint(100,999)}-{random.randint(10,99)}-{random.randint(1000,9999)}", "", "", "", "", "", ""])
        manifest.plant_false_positive(doc_id, "SSN-shaped order number", "ORD-7788 row", "Order reference number formatted like an SSN (###-##-####)")
        manifest.plant_false_positive(doc_id, "template placeholder row", "TEST USER row", "Placeholder/test record that must not be flagged as a real exposed person")

    wb.save(path)
    manifest.add_document(doc_id, filename, "xlsx", notes=f"{len(people_subset)} people in one export" + (" + FP traps" if include_fp_traps else ""))
    return path


def plant_csv(doc_id, people_subset, manifest):
    filename = f"{doc_id}.csv"
    path = os.path.join(OUT, filename)
    with open(path, "w") as f:
        f.write("full_name,dob,ssn_last4,card_number,medical_condition\n")
        for p in people_subset:
            # partial identifier edge case: only last-4 of SSN here
            last4 = p["ssn"][-4:]
            f.write(f'"{p["full_name"]}",{p["dob"]},{last4},{p["card_number"]},"{p["medical_condition"]}"\n')
            manifest.plant(doc_id, p["person_id"], "full_name", p["full_name"])
            manifest.plant(doc_id, p["person_id"], "dob", p["dob"])
            manifest.plant(doc_id, p["person_id"], "ssn", last4, edge_case="partial_identifier_last4", is_partial=True)
            manifest.plant(doc_id, p["person_id"], "card_number", p["card_number"])
            manifest.plant(doc_id, p["person_id"], "medical", p["medical_condition"])
    manifest.add_document(doc_id, filename, "csv", notes="Partial SSN (last-4) only — full SSN planted separately elsewhere")
    return path


# ----------------------------------------------------------------------
# EML with attachment
# ----------------------------------------------------------------------

def plant_eml(doc_id, person, manifest, edge_case=None, attach_doc_bytes=None, attach_name=None):
    filename = f"{doc_id}.eml"
    path = os.path.join(OUT, filename)

    template = random.choice(EMAIL_TEMPLATES)
    sender = _rand_signature()
    name = person["full_name"]
    if edge_case == "name_variant" and person["aliases"]:
        name = random.choice(person["aliases"])

    body = template.format(
        name=name, phone=person["phone"], email=person["email"], address=person["address"],
        dob=person["dob"], ssn=person["ssn"], card=person["card_number"],
        card_exp=person["card_exp"], sender=sender,
    )

    msg = EmailMessage()
    msg["Subject"] = random.choice(["Record Update", "Please review", "Account follow-up", "RE: Verification needed"])
    msg["From"] = f"{sender.splitlines()[0].lower().replace(' ', '.')}@datafactz-client.example.com"
    msg["To"] = "records@datafactz-client.example.com"
    msg.set_content(body)

    if attach_doc_bytes is not None:
        msg.add_attachment(attach_doc_bytes, maintype="application", subtype="octet-stream", filename=attach_name)

    with open(path, "wb") as f:
        f.write(bytes(msg))

    manifest.add_document(doc_id, filename, "eml", notes="Has attachment" if attach_doc_bytes else None)
    manifest.plant(doc_id, person["person_id"], "full_name", name, edge_case=edge_case)
    manifest.plant(doc_id, person["person_id"], "phone", person["phone"])
    manifest.plant(doc_id, person["person_id"], "email", person["email"])
    if "address" in template:
        manifest.plant(doc_id, person["person_id"], "home_address", person["address"])
    if "SSN" in template or "ssn" in template.lower():
        manifest.plant(doc_id, person["person_id"], "ssn", person["ssn"])
    if "card" in template:
        manifest.plant(doc_id, person["person_id"], "card_number", person["card_number"])
    if "DOB" in template:
        manifest.plant(doc_id, person["person_id"], "dob", person["dob"])
    return path


# ----------------------------------------------------------------------
# TXT (chat log style) and HTML
# ----------------------------------------------------------------------

def plant_txt(doc_id, person, manifest, edge_case=None):
    filename = f"{doc_id}.txt"
    path = os.path.join(OUT, filename)
    lines = [
        f"[09:14] agent: Can I get your full name and date of birth to pull up the account?",
        f"[09:14] customer: {person['full_name']}, born {person['dob']}",
        f"[09:15] agent: Thanks. And can you confirm the login email on file?",
        f"[09:15] customer: yes it's {person['email']}, username {person['username']}",
        f"[09:16] agent: One more — last 4 of the card ending in your account?",
        f"[09:16] customer: {person['card_number'][-4:]}",
    ]
    with open(path, "w") as f:
        f.write("\n".join(lines))
    manifest.add_document(doc_id, filename, "txt", notes="Support chat transcript")
    manifest.plant(doc_id, person["person_id"], "full_name", person["full_name"], edge_case=edge_case)
    manifest.plant(doc_id, person["person_id"], "dob", person["dob"])
    manifest.plant(doc_id, person["person_id"], "email", person["email"])
    manifest.plant(doc_id, person["person_id"], "login_credentials", person["username"])
    manifest.plant(doc_id, person["person_id"], "card_number", person["card_number"][-4:], is_partial=True, edge_case="partial_identifier_last4")
    return path


def plant_html(doc_id, person, manifest, edge_case=None):
    filename = f"{doc_id}.html"
    path = os.path.join(OUT, filename)
    html = f"""<html><body>
<h2>Customer Profile Export</h2>
<table border="1">
<tr><td>Name</td><td>{person['full_name']}</td></tr>
<tr><td>Address</td><td>{person['address']}</td></tr>
<tr><td>Phone</td><td>{person['phone']}</td></tr>
<tr><td>Login</td><td>{person['username']} / {person['password']}</td></tr>
</table>
</body></html>"""
    with open(path, "w") as f:
        f.write(html)
    manifest.add_document(doc_id, filename, "html")
    manifest.plant(doc_id, person["person_id"], "full_name", person["full_name"], edge_case=edge_case)
    manifest.plant(doc_id, person["person_id"], "home_address", person["address"])
    manifest.plant(doc_id, person["person_id"], "phone", person["phone"])
    manifest.plant(doc_id, person["person_id"], "login_credentials", f"{person['username']}/{person['password']}")
    return path


# ----------------------------------------------------------------------
# Problem files (no PII planted — these test triage/quarantine)
# ----------------------------------------------------------------------

def plant_zero_byte(doc_id, manifest):
    filename = f"{doc_id}.pdf"
    path = os.path.join(OUT, filename)
    open(path, "wb").close()
    manifest.add_document(doc_id, filename, "problem_zero_byte", quarantine_reason="Zero-byte file")
    return path


def plant_corrupt(doc_id, manifest):
    filename = f"{doc_id}.pdf"
    path = os.path.join(OUT, filename)
    with open(path, "wb") as f:
        f.write(os.urandom(400))  # garbage bytes, invalid PDF structure
    manifest.add_document(doc_id, filename, "problem_corrupt", quarantine_reason="Corrupt/unreadable file structure")
    return path


def plant_wrong_extension(doc_id, person, manifest):
    """A PNG saved with a .docx extension."""
    filename = f"{doc_id}.docx"
    path = os.path.join(OUT, filename)
    img = Image.new("RGB", (600, 300), "white")
    draw = ImageDraw.Draw(img)
    draw.text((20, 20), f"Record: {person['full_name']}", font=_get_font(20), fill="black")
    draw.text((20, 60), f"SSN: {person['ssn']}", font=_get_font(20), fill="black")
    img.save(path, "PNG")  # real format is PNG despite .docx name
    manifest.add_document(doc_id, filename, "problem_wrong_extension", quarantine_reason="File extension does not match actual content (PNG saved as .docx)")
    manifest.plant(doc_id, person["person_id"], "full_name", person["full_name"], edge_case="wrong_extension")
    manifest.plant(doc_id, person["person_id"], "ssn", person["ssn"], edge_case="wrong_extension")
    return path


def plant_password_protected(doc_id, person, manifest):
    from pypdf import PdfWriter
    filename = f"{doc_id}.pdf"
    tmp_path = os.path.join(OUT, f"_tmp_{doc_id}.pdf")
    path = os.path.join(OUT, filename)

    c = pdfcanvas.Canvas(tmp_path, pagesize=letter)
    c.drawString(72, 700, f"Confidential record for {person['full_name']}")
    c.drawString(72, 680, f"SSN: {person['ssn']}")
    c.save()

    writer = PdfWriter()
    from pypdf import PdfReader
    reader = PdfReader(tmp_path)
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("breach2026")
    with open(path, "wb") as f:
        writer.write(f)
    os.remove(tmp_path)

    manifest.add_document(doc_id, filename, "problem_password_protected", quarantine_reason="Password-protected — cannot parse without credential")
    manifest.plant(doc_id, person["person_id"], "full_name", person["full_name"], edge_case="password_protected")
    manifest.plant(doc_id, person["person_id"], "ssn", person["ssn"], edge_case="password_protected")
    return path
